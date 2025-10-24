"""
DOCX to Markdown Converter
Converts Microsoft Word documents to markdown format preserving structure
"""
import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any
import docx
import docx2txt


def convert(file_path: Path, output_dir: Path) -> Dict[str, Any]:
    """
    Convert DOCX file to markdown format
    
    Args:
        file_path: Path to the input DOCX file
        output_dir: Directory where output should be written
    
    Returns:
        Dictionary with conversion result
    """
    try:
        # Create docx output directory if it doesn't exist
        docx_dir = output_dir / "docx"
        docx_dir.mkdir(exist_ok=True)
        
        # Read the DOCX file
        doc = docx.Document(str(file_path))
        
        # Build markdown content
        markdown_lines = []
        
        # Add metadata header
        markdown_lines.extend([
            "---",
            f"source: {file_path.relative_to(file_path.parent)}",
            f"converted_at: {datetime.utcnow().isoformat()}Z",
            "converter: docx_to_md",
            "---",
            ""
        ])
        
        # Process paragraphs and preserve structure
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Determine paragraph style
            style = paragraph.style.name.lower()
            
            if style.startswith('heading'):
                # Convert heading levels (Heading 1 -> #, Heading 2 -> ##, etc.)
                import re
                match = re.search(r'\d+', style)
                if match:
                    level = int(match.group())
                else:
                    level = 1
                level = max(1, min(6, level))  # Limit to h1-h6
                markdown_lines.append(f"{'#' * level} {text}")
            else:
                # Handle bullet lists and regular text
                if paragraph.style.name.lower() in ['list paragraph', 'bullet']:
                    markdown_lines.append(f"- {text}")
                else:
                    markdown_lines.append(text)
            
            markdown_lines.append("")  # Add blank line after each paragraph
        
        # Process tables if any
        for table in doc.tables:
            if table.rows:
                markdown_lines.append("### Table")
                # Add table header
                header_cells = table.rows[0].cells
                header_line = "| " + " | ".join([cell.text.strip() for cell in header_cells]) + " |"
                separator_line = "| " + " | ".join(["---"] * len(header_cells)) + " |"
                markdown_lines.extend([header_line, separator_line])
                
                # Add table rows
                for row in table.rows[1:]:
                    row_line = "| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |"
                    markdown_lines.append(row_line)
                
                markdown_lines.append("")  # Blank line after table
        
        # Join all lines
        markdown_content = "\n".join(markdown_lines).strip()
        
        # Write output file
        output_filename = file_path.stem + ".md"
        output_file = docx_dir / output_filename
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        return {
            'status': 'success',
            'output': str(output_file.relative_to(output_dir.parent)),
            'converter': 'docx_to_md',
            'metadata': {
                'metadata': {
                    'sections': len(doc.sections),
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables)
                }                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
        }
        
    except Exception as e:
        return {
            'status': 'failed',
            'error': str(e),
            'converter': 'docx_to_md'
        }


# Alternative implementation using docx2txt for simpler text extraction
def convert_simple(file_path: Path, output_dir: Path) -> Dict[str, Any]:
    """
    Simple DOCX to text conversion using docx2txt
    """
    try:
        # Create docx output directory if it doesn't exist
        docx_dir = output_dir / "docx"
        docx_dir.mkdir(exist_ok=True)
        
        # Extract text using docx2txt
        text = docx2txt.process(str(file_path))
        
        # Add metadata header
        markdown_content = f"""---
source: {file_path.relative_to(file_path.parent)}
converted_at: {datetime.utcnow().isoformat()}Z
converter: docx_to_md
---

{text}
"""
        
        # Write output file
        output_filename = file_path.stem + ".md"
        output_file = docx_dir / output_filename
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # Count paragraphs roughly
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        return {
            'status': 'success',
            'output': str(output_file.relative_to(output_dir.parent)),
            'converter': 'docx_to_md',
            'metadata': {
                'paragraphs': paragraphs
            }
        }
        
    except Exception as e:
        return {
            'status': 'failed',
            'error': str(e),
            'converter': 'docx_to_md'
        }
